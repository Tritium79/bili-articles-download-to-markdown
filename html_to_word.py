import argparse
import os
import re
import requests
import docx
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches


def download_image(url, save_dir):
    """
    下载图片并保存到本地

    参数:
    url (str): 图片URL
    save_dir (str): 保存目录

    返回:
    str: 本地图片路径，如果下载失败则返回None
    """
    try:
        # 处理相对URL
        if url.startswith('//'):
            url = 'https:' + url
        elif not url.startswith(('http://', 'https://')):
            # 如果是完全相对路径，这里可以添加基础URL
            # 对于B站，我们可以使用https://i0.hdslb.com作为基础
            url = 'https://i0.hdslb.com' + url

        # 获取图片文件名
        filename = os.path.basename(url.split('?')[0])
        save_path = os.path.join(save_dir, filename)

        # 确保保存目录存在
        os.makedirs(save_dir, exist_ok=True)

        # 下载图片
        response = requests.get(url, timeout=10)
        response.raise_for_status()

        # 保存图片
        with open(save_path, 'wb') as f:
            f.write(response.content)

        print(f"图片已下载到: {save_path}")
        return save_path

    except Exception as e:
        print(f"下载图片失败: {e}")
        return None


def html_to_word(input_file, output_file):
    """
    将HTML文件转换为Word文档

    参数:
    input_file (str): 输入的HTML文件路径
    output_file (str): 输出的Word文件路径

    返回:
    bool: 成功返回True，失败返回False
    """
    try:
        # 打印调试信息
        print(f"开始转换HTML到Word: {input_file} -> {output_file}")
        print(f"当前工作目录: {os.getcwd()}")
        print(f"输出文件绝对路径: {os.path.abspath(output_file)}")

        # 创建Word文档
        doc = Document()
        
        # 设置默认字体为雅黑
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        # 设置中文字体
        style.element.rPr.rFonts.set(docx.oxml.shared.qn('w:eastAsia'), 'Microsoft YaHei')

        # 读取HTML文件
        if not os.path.exists(input_file):
            print(f"错误: 输入文件不存在 - {input_file}")
            return False

        with open(input_file, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # 过滤无法用gbk编码的字符
        html_content = html_content.encode('gbk', 'ignore').decode('gbk')

        # 解析HTML
        soup = BeautifulSoup(html_content, 'html.parser')

        # 提取网页标题
        title = "output"
        if soup.title:
            title = soup.title.string.strip()
            # 移除非法文件名字符
            title = re.sub(r'[\\/:*?"<>|]', '_', title)
        print(f"提取到网页标题: {title}")

        # 检查body标签是否存在
        if not soup.body:
            print("警告: HTML中未找到body标签，将使用整个文档进行处理")
            # 使用整个文档代替body
            content_to_process = soup
        else:
            content_to_process = soup.body

        # 打印content_to_process的类型和内容
        print(f"content_to_process类型: {type(content_to_process)}")
        print(f"content_to_process内容预览: {str(content_to_process)[:100]}...")

        # 创建图片保存目录
        img_dir = os.path.join(os.path.dirname(output_file), 'images')
        os.makedirs(img_dir, exist_ok=True)
        print(f"图片保存目录: {img_dir}")

        # 遍历HTML内容
        element_count = 0

        def process_element(element):
            nonlocal element_count
            element_count += 1
            print(f"处理元素: 类型={type(element)}, 名称={element.name}")
            
            # 处理文本节点
            if element.name is None:
                text = str(element).strip()
                if text:
                    # 替换多个空白字符为单个空格
                    import re
                    text = re.sub(r'\s+', ' ', text)
                    # 过滤无法编码的字符
                    text = text.encode('gbk', 'ignore').decode('gbk')
                    # 检查是否已有段落，如果没有则创建一个
                    if not doc.paragraphs or doc.paragraphs[-1].text.strip():
                        doc.add_paragraph(text)
                    else:
                        doc.paragraphs[-1].add_run(text)
                    print(f"添加文本: {text[:30]}...")
            elif element.name == 'p':
                # 处理段落
                text = element.get_text(strip=False)
                if text.strip():
                    # 替换多个空白字符为单个空格
                    import re
                    text = re.sub(r'\s+', ' ', text)
                    para = doc.add_paragraph(text)
                    # 设置段落格式
                    para.paragraph_format.space_after = Inches(0.1)
                    para.paragraph_format.first_line_indent = Inches(0.25)
                    print(f"添加段落: {text[:30]}...")
            elif element.name == 'br':
                # 处理换行标签
                doc.add_paragraph()
                print("添加换行")
            elif element.name == 'a':
                # 处理超链接
                link_text = element.get_text(strip=True)
                link_url = element.get('href')
                if link_text and link_url:
                    # 检查是否已有段落，如果没有则创建一个
                    if not doc.paragraphs or doc.paragraphs[-1].text.strip():
                        para = doc.add_paragraph()
                    else:
                        para = doc.paragraphs[-1]
                    # 添加超链接
                    run = para.add_run(link_text)
                    run.font.underline = True
                    run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
                    # 这里只是设置了视觉效果，实际超链接功能需要使用hyperlink
                    # 但python-docx不直接支持，需要使用更复杂的方法
                    print(f"添加超链接: {link_text} -> {link_url}")
            elif element.name == 'img':
                # 处理图片
                img_url = element.get('src')
                if img_url:
                    print(f"找到图片URL: {img_url}")
                    img_path = download_image(img_url, img_dir)
                    if img_path:
                        # 计算页面宽度和高度的一半（假设标准A4纸，宽度约6.5英寸，高度约9英寸）
                        max_width = Inches(3.25)  # 页面宽度的一半
                        max_height = Inches(4.5)  # 页面高度的一半
                        
                        # 添加图片并设置最大尺寸
                        pic = doc.add_picture(img_path)
                        # 获取原始图片尺寸
                        orig_width, orig_height = pic.width, pic.height
                        
                        # 计算调整后的尺寸，保持 aspect ratio
                        width_scale = max_width / orig_width
                        height_scale = max_height / orig_height
                        scale = min(width_scale, height_scale)
                        
                        # 如果图片不需要缩放，则不改变尺寸
                        if scale < 1:
                            pic.width = int(orig_width * scale)
                            pic.height = int(orig_height * scale)
                        
                        print(f"添加图片: {img_path}, 调整后尺寸: {pic.width}x{pic.height}")
            elif element.name in ['div', 'span', 'section', 'article']:
                # 递归处理容器元素
                for child in element.children:
                    process_element(child)
            # 可以根据需要添加更多标签处理

        # 开始递归处理
        if hasattr(content_to_process, 'children'):
            for child in content_to_process.children:
                process_element(child)
        else:
            process_element(content_to_process)

        print(f"共处理 {element_count} 个元素")

        # 保存Word文档
        # 确保输出目录存在
        output_dir = os.path.dirname(output_file)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
            print(f"创建输出目录: {output_dir}")
        
        try:
            doc.save(output_file)
            # 验证文件是否已保存
            if os.path.exists(output_file):
                print(f"成功转换HTML为Word并保存到: {output_file}")
                print(f"输出文件大小: {os.path.getsize(output_file)} 字节")
                return True
            else:
                print(f"警告: 保存文件后检查失败，文件不存在: {output_file}")
                return False
        except Exception as e:
            print(f"保存Word文档失败: {e}")
            return False

    except Exception as e:
        print(f"转换HTML为Word失败: {e}")
        return False


if __name__ == "__main__":
    # 解析命令行参数
    parser = argparse.ArgumentParser(description='将HTML文件转换为Word文档')
    parser.add_argument('input', help='输入的HTML文件路径')
    parser.add_argument('-o', '--output', help='输出的Word文件名(默认: 网页标题.docx)')

    args = parser.parse_args()

    # 确定输出文件名
    output_file = args.output
    if not output_file:
        # 如果未指定输出文件名，则使用网页标题
        with open(args.input, 'r', encoding='utf-8') as f:
            html_content = f.read()
        soup = BeautifulSoup(html_content, 'html.parser')
        title = "output"
        if soup.title:
            title = soup.title.string.strip()
            # 移除非法文件名字符
            title = re.sub(r'[\\/:*?"<>|]', '_', title)
        output_file = f"{title}.docx"
        print(f"使用网页标题作为输出文件名: {output_file}")

    # 调用函数转换HTML为Word
    html_to_word(args.input, output_file)